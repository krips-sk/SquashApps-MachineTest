# -*- coding: utf-8 -*-
# Name: service.py
# Description: To generate Workshop Expense Report
# Author: Mycura
# Created: 2020.12.30
# Copyright: © 2020 . All Rights Reserved.
# Change History:
#                |2020.12.30
import calendar
import logging
import re
import socket
import uuid
from datetime import datetime
from dateutil.parser import parse

import psycopg2
from docx.api import Document
import xlrd
from pandas import *
from ems_config import pg_username, pg_password, pg_host, pg_port, pg_database
import sys
import click
from xlrd.timemachine import xrange

logging.basicConfig(
    handlers=[logging.FileHandler('corrective.log', 'w', 'utf-8')],
    format='%(levelname)s: %(message)s',
    datefmt='%m-%d %H:%M',
    level=logging.INFO  # CRITICAL ERROR WARNING  INFO    DEBUG    NOTSET
)
number_empty = ['nil', 'nan', '', 'null']

unwantted_chars = ['nil', '(rm)']


def main():
    try:
        print(" SOCKET CONNECTION STARTS ")
        # Create a socket (SOCK_STREAM means a TCP socket)
        # sock = socket.socket(socket.AF_INET, socket.SOCK_STREAM)
        # Connect to server and send data
        # sock.bind((HOST, PORT))
        # process_file('connection', 'c_date', 'user_id', 'full_path')
        generate_workshop_corrective_report()
    except Exception as e:
        print("Main Error" + str(e))
        logging.error(e)
    finally:
        print(" CONNECTION CLOSED ")
        # sock.close()


def is_number(s):
    """ Returns True is string is a number. """
    return s.replace('.', '', 1).isdigit()


def get_product_details(table):
    try:
        # Data will be a list of rows represented as dictionaries
        # containing each row's data.
        data = []

        keys = None
        data = []
        old_item = ''
        for i, row in enumerate(table.rows):
            text = (cell.text for cell in row.cells)

            # Establish the mapping based on the first row
            # headers; these will become the keys of our dictionary
            if i == 0:
                text = (remove_useless_char(cell.text.strip().lower()) for cell in row.cells)
                keys = tuple(text)
                continue

            # Construct a dictionary for this row, mapping
            # keys to values for this row
            row_data = dict(zip(keys, text))
            data.append(row_data)
        return data
    except Exception as e:
        print("Main Error" + str(e))
        logging.error(e)


def get_product_info(table):
    try:
        # finding the sub table.
        for i, row in enumerate(table.rows):
            for col_index, c_item in enumerate(row.cells):
                if c_item.tables:
                    return get_product_details(c_item.tables[0])
    except Exception as e:
        print("Main Error" + str(e))
        logging.error(e)


def get_supplier_info(table):
    try:
        # Data will be a list of rows represented as dictionaries
        # containing each row's data.
        data = []

        keys, row_data = None, {'name': '', 'department': '', 'car_no': '', 'date': ''}
        old_item = ''
        for i, row in enumerate(table.rows):
            for col_index, c_item in enumerate(row.cells):
                item = str(c_item.text).strip().lower()
                if item:
                    data = remove_prefix(c_item.text, ':').strip()
                    data = re.sub('\s+', ' ', data)
                    item = item.replace('car no.', 'car_no')
                    if item in ['name', 'department', 'car_no', 'date']:
                        old_item = item
                    else:
                        row_data[old_item] = data
        return row_data
    except Exception as e:
        print("Main Error" + str(e))
        logging.error(e)


def get_category_info(table):
    try:
        keys, row_data = None, {'quality': 0, 'environment': 0, 'safety': 0, 'health': 0}
        old_item = ''
        for i, row in enumerate(table.rows):
            for col_index, c_item in enumerate(row.cells):
                item = str(c_item.text).lower().replace('(q)', '') \
                    .replace('(e)', '').replace('(s)', '').replace('(h)', '').strip()
                if item:
                    if item in ['quality', 'environment', 'safety', 'health']:
                        old_item = item
                    else:
                        row_data[old_item] = 1
        return row_data
    except Exception as e:
        print("Main Error" + str(e))
        logging.error(e)


def process_file(connection, c_date, user_id, full_path):
    try:
        full_path = 'C:\\Users\\Thiru_office\\Documents\\EMS_malaka\\RETURNED GOODS FORM- format.docx'
        document = Document(full_path)
        supplier_info = {}
        product_info = []
        category_info = {}
        for table in document.tables:
            if supplier_info:
                category_info = get_category_info(table)
                product_info = get_product_info(table)
            else:
                supplier_info = get_supplier_info(table)

            if supplier_info and product_info:
                # need to process and save details to database
                save_details_to_db(connection, c_date, user_id, supplier_info, product_info, category_info)
                supplier_info = {}
                product_info = {}


    except Exception as e:
        print("Main Error" + str(e))
        logging.error(e)


#
#
#
def generate_workshop_corrective_report():
    try:
        # Logic
        # 1. Get the list of pending
        # 2 do until remaining
        # 3. process
        # 4. check remaining
        connection = psycopg2.connect(user=pg_username,
                                      password=pg_password,
                                      host=pg_host,
                                      port=pg_port,
                                      database=pg_database)

        try:
            connection.autocommit = True
            print("Connection success - Job Started...")
            logging.error("job started")
            query = connection.cursor()
            pending_query = """
            SELECT report_id, report_date, created_by, full_path	FROM public."tblReportTracker" 
            WHERE status=0 AND report_type=3 AND "isActive"=1; 
            """
            query.execute(pending_query)
            pending_data = query.fetchall()
            for item in pending_data:
                c_date = item[1]
                user_id = item[2]
                report_id = item[0]
                full_path = item[3]
                logging.error("picked report : {report_id}".format(report_id=report_id))
                update_report_status(connection, report_id, 1)
                process_file(connection, c_date, user_id, full_path)
                update_report_status(connection, report_id, 4)
            connection.commit()
            query.close()
            print("Job Completed")
        except Exception as e:
            logging.error(e)
            print("Generate Workshop parts Error : " + str(e))
        connection.close()

    except Exception as e:
        logging.error(e)
        print("Generate Workshop parts Error : " + str(e))


def save_details_to_db(connection, c_date, user_id, supplier_info, product_info, category_info):
    try:
        # Get the department id
        # Get the supplier id
        # Convert report date
        # save head
        # iterate over products -> skip total
        supplier_id = get_supplier_id(connection, supplier_info['name'])
        department_id = get_department_id(connection, supplier_info['department'])
        report_date_obj = convertdmy_to_date2(supplier_info['date'])
        report_id = get_report_head_id(connection, supplier_id, department_id, report_date_obj, supplier_info['car_no'],
                                       category_info)
        query = connection.cursor()
        for product in product_info:
            if product['description']:
                if 'part_no' not in product:
                    product['part_no'] = ''
                product['price'] = remove_null(product['price'])
                product['total'] = remove_null(product['total'])
                report_insert_query = """
                    INSERT INTO "tblReportCorrectiveDetail" ( head_id,description,part_no,qty,price,
                    total,
                     created_by, created_date, "isActive" ) VALUES 
                                          ( {head_id},'{description}','{part_no}',{qty},{price},
                                          {total},
                                          {user_id},'{created_date}',1)
                    """
                report_insert_query = report_insert_query.format(head_id=report_id,
                                                                 description=product['description'],
                                                                 part_no=product['part_no'],
                                                                 qty=product['qty'],
                                                                 price=product['price'],
                                                                 total=product['total'],
                                                                 user_id=1,
                                                                 created_date=datetime.utcnow())

                query.execute(report_insert_query)
        query.close()
    except Exception as e:
        logging.error(e)
        print("Save details Error :" + str(e))


def get_report_head_id(connection, supplier_id, department_id, report_date_obj, car_no, category_info):
    try:
        head_id = get_exists_head(connection, supplier_id, department_id, report_date_obj, car_no)
        if head_id > 0:
            return head_id
        else:
            query = connection.cursor()
            report_insert_query = """
                            INSERT INTO "tblReportCorrectiveHead" ( report_date, car_no,supplier_id,department_id,
                             quality,environment,safety,health, created_by, created_date, "isActive" ) VALUES 
                                                  ( '{report_date}','{car_no}',{supplier_id},{department_id},
                                                  {quality},{environment},{safety},{health},
                                                  {user_id},'{created_date}',1)
                            """
            report_insert_query = report_insert_query.format(report_date=report_date_obj,
                                                             car_no=car_no,
                                                             supplier_id=supplier_id,
                                                             department_id=department_id,
                                                             quality=category_info['quality'],
                                                             environment=category_info['environment'],
                                                             safety=category_info['safety'],
                                                             health=category_info['health'],
                                                             user_id=1,
                                                             created_date=datetime.utcnow())

            query.execute(report_insert_query)
            query.close()
            head_id = get_exists_head(connection, supplier_id, department_id, report_date_obj, car_no)
            return head_id
    except Exception as e:
        logging.error(e)
        print("Generate Workshop parts Error : " + str(e))


def get_exists_head(connection, supplier_id, department_id, report_date_obj, car_no):
    try:
        head_id = 0
        query = connection.cursor()
        exists_query = """ SELECT report_id FROM "tblReportCorrectiveHead" WHERE "isActive"=1 AND car_no ilike '{car_no}' 
        AND supplier_id={supplier_id} AND department_id={department_id} AND report_date='{report_date}'; 
                        """
        exists_query = exists_query.format(car_no=car_no,
                                           supplier_id=supplier_id,
                                           department_id=department_id,
                                           report_date=report_date_obj
                                           )
        query.execute(exists_query)
        exists_data = query.fetchone()
        if exists_data:
            head_id = exists_data[0]
        query.close()
        return head_id
    except Exception as e:
        logging.error(e)
        print("Generate Workshop parts Error : " + str(e))
        return 0


def get_exists_supplier(connection, supplier):
    try:
        supplier_id = 0
        query = connection.cursor()
        exists_query = """ SELECT id FROM "tblSupplier" WHERE "isactive"=1 AND supplier_name ilike '{supplier}' ; 
                        """
        exists_query = exists_query.format(supplier=supplier)
        query.execute(exists_query)
        exists_data = query.fetchone()
        if exists_data:
            supplier_id = exists_data[0]
        query.close()
        return supplier_id
    except Exception as e:
        logging.error(e)
        print("Generate Workshop parts Error : " + str(e))
        return 0


def get_supplier_id(connection, supplier):
    try:
        supplier_id = get_exists_supplier(connection, supplier)
        if supplier_id > 0:
            return supplier_id
        else:
            public_id = str(uuid.uuid4())

            query = connection.cursor()
            report_insert_query = """
                            INSERT INTO "tblSupplier" ( public_id, supplier_name,
                             created_by, created_date, "isactive" ) VALUES 
                                                  ( '{public_id}','{supplier}',
                                                  {user_id},'{created_date}',1)
                            """
            report_insert_query = report_insert_query.format(public_id=public_id,
                                                             supplier=supplier,
                                                             user_id=1,
                                                             created_date=datetime.utcnow())

            query.execute(report_insert_query)
            query.close()
            supplier_id = get_exists_supplier(connection, supplier)
            return supplier_id
    except Exception as e:
        logging.error(e)
        print("Generate Workshop parts Error : " + str(e))


def get_exists_department(connection, department):
    try:
        department_id = 0
        query = connection.cursor()
        exists_query = """ SELECT id FROM "tbldepartment" WHERE "isactive"=1 AND department_name ilike '{department}' ; 
                        """
        exists_query = exists_query.format(department=department)
        query.execute(exists_query)
        exists_data = query.fetchone()
        if exists_data:
            department_id = exists_data[0]
        query.close()
        return department_id
    except Exception as e:
        logging.error(e)
        print("Generate Workshop parts Error : " + str(e))
        return 0


def get_department_id(connection, department):
    try:
        department_id = get_exists_department(connection, department)
        if department_id > 0:
            return department_id
        else:
            public_id = str(uuid.uuid4())

            query = connection.cursor()
            report_insert_query = """
                            INSERT INTO "tbldepartment" ( public_id, department_name,
                             created_by, created_date, "isactive" ) VALUES 
                                                  ( '{public_id}','{department}',
                                                  {user_id},'{created_date}',1)
                            """
            report_insert_query = report_insert_query.format(public_id=public_id,
                                                             department=department,
                                                             user_id=1,
                                                             created_date=datetime.utcnow())

            query.execute(report_insert_query)
            query.close()
            department_id = get_exists_supplier(connection, department)
            return department_id
    except Exception as e:
        logging.error(e)
        print("Generate Workshop parts Error : " + str(e))


def update_report_status(connection, report_id, status):
    try:
        query = connection.cursor()
        update_report_query = """UPDATE public."tblReportTracker" SET status={status}, updated_date='{updated_date}',
         updated_by=1 WHERE report_id={report_id};
        """
        update_report_query = update_report_query.format(updated_date=datetime.utcnow(),
                                                         status=status,
                                                         report_id=report_id)
        query.execute(update_report_query)
        query.close()
        return 1
    except Exception as e:
        logging.error(e)
        print("Update Status Error : " + str(e))


def remove_null(c_item):
    return '0' if c_item in number_empty else c_item


def remove_prefix(text, prefix):
    if text.startswith(prefix):
        return text[len(prefix):]
    return text  # or whatever


def remove_useless_char(word, is_header_replace=True):
    for remove_char in unwantted_chars:
        word = word.replace(remove_char, '')
        # replace headers
    if is_header_replace:
        word = word.replace('part number', 'part_no')
    word = word.strip()
    return word


def convertdmy_to_date2(strdate):
    try:
        strdate = strdate.replace(".", "-").replace("/", "-")
        return parse(strdate)
    except Exception as e:
        logging.exception(e)
        print(e)
        print('Handling run-time error:')


if __name__ == "__main__":
    main()
